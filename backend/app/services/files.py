"""
File ingestion (user attachments) and document generation (assistant
outputs). Text-heavy formats are parsed into plain text that can be
embedded in Claude's prompt; images are passed through as base64 so
Claude can see them natively via vision.
"""

from __future__ import annotations

import base64
import io
import logging
import re
from typing import Any, Dict, List, Optional, Tuple

from docx import Document
from docx.shared import Pt, RGBColor
from pypdf import PdfReader

logger = logging.getLogger(__name__)


# ── Bytea wire-format decoding ───────────────────────────────────────
def decode_bytea(value) -> bytes:
    """
    Supabase returns `bytea` values as strings over PostgREST. Postgres
    uses two wire formats:
      - hex: "\\x" followed by lowercase hex digits (the default)
      - escape: backslash-escaped bytes (legacy)
    We also handle the cases where PostgREST has already given us raw
    bytes, or base64 (some setups do that).
    """
    if isinstance(value, (bytes, bytearray)):
        return bytes(value)
    if not isinstance(value, str):
        raise ValueError(f"unexpected bytea wire type: {type(value)}")

    if value.startswith("\\x") or value.startswith("\\\\x"):
        hex_body = value[2:] if value.startswith("\\x") else value[3:]
        return bytes.fromhex(hex_body)

    try:
        return base64.b64decode(value, validate=True)
    except Exception as e:
        raise ValueError(f"failed to decode bytea: {e}")


def encode_bytes_for_bytea(data: bytes) -> str:
    """
    Postgres `bytea` accepts the hex-escape format ``\\x<hex>``, which is
    what PostgREST expects when writing binary columns over JSON.
    """
    return "\\x" + data.hex()


# ── Limits ───────────────────────────────────────────────────────────
MAX_FILE_BYTES = 10 * 1024 * 1024       # 10 MB / file
MAX_FILES_PER_MESSAGE = 3

# Pre-declared so we can use a single source of truth in the route.
ALLOWED_MIME_TYPES = {
    "application/pdf",
    "text/plain",
    "text/markdown",
    "image/png",
    "image/jpeg",
    "image/jpg",
}

IMAGE_MIME_TYPES = {"image/png", "image/jpeg", "image/jpg"}

# Cap the amount of extracted text we feed into Claude per document so
# a 200-page PDF doesn't blow the context window. ~50k chars ≈ 12k tokens.
MAX_EXTRACTED_CHARS = 50_000


# ── Ingestion ────────────────────────────────────────────────────────
def extract_pdf_text(data: bytes) -> str:
    """Pull plain text out of a PDF. Best-effort — returns empty on parse errors."""
    try:
        reader = PdfReader(io.BytesIO(data))
        chunks: List[str] = []
        for page in reader.pages:
            try:
                chunks.append(page.extract_text() or "")
            except Exception as e:  # malformed page
                logger.warning(f"PDF page extract failed: {e}")
        return "\n\n".join(c for c in chunks if c).strip()
    except Exception as e:
        logger.warning(f"PDF parse failed: {e}")
        return ""


def build_claude_content_blocks(
    user_message: str,
    files: List[Tuple[str, str, bytes]],
    *,
    file_ids: Optional[List[Optional[str]]] = None,
) -> List[Dict[str, Any]]:
    """
    Convert a user message plus attachment payloads into Claude's
    content-block format.

    `files` is a list of (filename, mime_type, bytes). `file_ids`, if
    provided, is a parallel list of Anthropic-Files-API IDs (or None)
    for each attachment. When a file_id is present AND the mime is
    supported by the Files API (PDFs, images), we reference the upload
    by id instead of inlining base64 — which on follow-up turns saves
    both network bandwidth and ~15-25k input tokens per PDF.

    Text-like files (.txt / .md) always use inline text extraction —
    the Files API adds no benefit there.
    """
    blocks: List[Dict[str, Any]] = []
    if file_ids is None:
        file_ids = [None] * len(files)

    text_preludes: List[str] = []
    for (filename, mime, data), file_id in zip(files, file_ids):
        # Image: prefer file_id reference on follow-up turns.
        if mime in IMAGE_MIME_TYPES:
            media_type = "image/jpeg" if mime == "image/jpg" else mime
            if file_id:
                blocks.append({
                    "type": "image",
                    "source": {"type": "file", "file_id": file_id},
                })
            else:
                blocks.append({
                    "type": "image",
                    "source": {
                        "type": "base64",
                        "media_type": media_type,
                        "data": base64.b64encode(data).decode("ascii"),
                    },
                })
            continue

        # PDF: prefer a native document block via file_id. Text
        # extraction is still the fallback when upload isn't available
        # (e.g. historical attachments pre-migration).
        if mime == "application/pdf" and file_id:
            blocks.append({
                "type": "document",
                "source": {"type": "file", "file_id": file_id},
                "title": filename,
                "citations": {"enabled": True},
            })
            continue

        if mime == "application/pdf":
            text = extract_pdf_text(data)
        elif mime in ("text/plain", "text/markdown"):
            try:
                text = data.decode("utf-8", errors="replace")
            except Exception as e:
                logger.warning(f"Text decode failed for {filename}: {e}")
                text = ""
        else:
            text = ""

        if not text:
            text_preludes.append(
                f"[Attached: {filename} — could not extract text]"
            )
            continue

        if len(text) > MAX_EXTRACTED_CHARS:
            text = text[:MAX_EXTRACTED_CHARS] + "\n…[truncated]"

        text_preludes.append(
            f"[Attached: {filename}]\n---\n{text}\n---"
        )

    combined_text = ""
    if text_preludes:
        combined_text = "\n\n".join(text_preludes) + "\n\n"
    combined_text += user_message

    blocks.append({"type": "text", "text": combined_text})

    return blocks


# ── Generation (Markdown → DOCX) ─────────────────────────────────────
_MD_HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
_MD_BULLET_RE = re.compile(r"^\s*[-*]\s+(.*)$")
_MD_NUMBERED_RE = re.compile(r"^\s*\d+\.\s+(.*)$")
_MD_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
_MD_ITALIC_RE = re.compile(r"(?<!\*)\*([^*\n]+)\*(?!\*)")
_MD_CODE_INLINE_RE = re.compile(r"`([^`\n]+)`")


def _add_rich_runs(paragraph, text: str) -> None:
    """
    Minimal Markdown inline parser for **bold**, *italic*, and `code`.
    Renders into docx runs on the given paragraph.
    """
    # Tokenise by splitting on each pattern in turn. Simple, not perfect,
    # but covers the common cases Claude emits.
    cursor = 0
    pattern = re.compile(
        r"(\*\*.+?\*\*|\*[^*\n]+\*|`[^`\n]+`)", re.DOTALL
    )
    for match in pattern.finditer(text):
        if match.start() > cursor:
            paragraph.add_run(text[cursor : match.start()])
        token = match.group(0)
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("*") and token.endswith("*"):
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        elif token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Menlo"
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        cursor = match.end()
    if cursor < len(text):
        paragraph.add_run(text[cursor:])


def markdown_to_docx_bytes(markdown: str, title: str | None = None) -> bytes:
    """
    Convert a Markdown string into a DOCX file (bytes). Supports
    headings, paragraphs, unordered/ordered lists, bold/italic/code.
    """
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    if title:
        doc.add_heading(title, level=0)

    lines = markdown.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]

        if not line.strip():
            i += 1
            continue

        heading_match = _MD_HEADING_RE.match(line)
        if heading_match:
            level = min(len(heading_match.group(1)), 4)
            doc.add_heading(heading_match.group(2).strip(), level=level)
            i += 1
            continue

        bullet_match = _MD_BULLET_RE.match(line)
        if bullet_match:
            while i < len(lines) and _MD_BULLET_RE.match(lines[i]):
                p = doc.add_paragraph(style="List Bullet")
                _add_rich_runs(p, _MD_BULLET_RE.match(lines[i]).group(1).strip())
                i += 1
            continue

        numbered_match = _MD_NUMBERED_RE.match(line)
        if numbered_match:
            while i < len(lines) and _MD_NUMBERED_RE.match(lines[i]):
                p = doc.add_paragraph(style="List Number")
                _add_rich_runs(p, _MD_NUMBERED_RE.match(lines[i]).group(1).strip())
                i += 1
            continue

        # Paragraph — consume until blank line or block break
        para_lines: List[str] = []
        while i < len(lines) and lines[i].strip() and not (
            _MD_HEADING_RE.match(lines[i])
            or _MD_BULLET_RE.match(lines[i])
            or _MD_NUMBERED_RE.match(lines[i])
        ):
            para_lines.append(lines[i].rstrip())
            i += 1
        p = doc.add_paragraph()
        _add_rich_runs(p, " ".join(para_lines))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def markdown_to_bytes(markdown: str) -> bytes:
    return markdown.encode("utf-8")
